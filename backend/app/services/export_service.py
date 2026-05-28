import os
from sqlalchemy.orm import Session
from typing import Optional, List
from app.config import settings
from app.repositories.project_repository import ProjectRepository
from app.repositories.citation_repository import CitationRepository

class ExportService:
    @staticmethod
    def compile_document(db: Session, project_id: int, citation_style: str = "ieee") -> dict:
        """
        Gathers all sections and formats them, assembling the citations bibliography.
        """
        project = ProjectRepository.get_by_id(db, project_id)
        if not project:
            return {"title": "Untitled", "content": "", "bibliography": ""}

        # Ordering of academic sections
        section_order = ["Abstract", "Introduction", "Literature Review", "Methodology", "Results & Discussion", "Conclusion"]
        
        compiled_text = []
        for sec_name in section_order:
            section = ProjectRepository.get_latest_section(db, project_id, sec_name)
            if section and section.content:
                compiled_text.append(f"## {sec_name}\n\n{section.content}")

        # Build bibliography
        citations = CitationRepository.get_by_project(db, project_id)
        bib_items = []
        
        if citations:
            if citation_style.lower() == "ieee":
                for idx, cite in enumerate(citations, 1):
                    # Format as: [1] Authors, "Title," venue, year.
                    bib_items.append(f"[{idx}] {cite.ieee or (cite.authors + ', ' + cite.title + ', ' + str(cite.year))}")
            else: # APA style
                # Sort alphabetically by author name
                sorted_citations = sorted(citations, key=lambda x: x.authors or "")
                for cite in sorted_citations:
                    bib_items.append(cite.apa or f"{cite.authors} ({cite.year}). {cite.title}.")

        bibliography = "## References\n\n" + "\n\n".join(bib_items) if bib_items else ""
        
        return {
            "title": project.title,
            "content": "\n\n".join(compiled_text),
            "bibliography": bibliography
        }

    @staticmethod
    def export_as_docx(db: Session, project_id: int, citation_style: str = "ieee") -> str:
        """
        Creates a Word DOCX file and returns the local file path.
        """
        doc_data = ExportService.compile_document(db, project_id, citation_style)
        filename = f"Project_{project_id}_Draft.docx"
        file_path = os.path.join(settings.DOCS_DIR, filename)

        try:
            from docx import Document
            doc = Document()
            doc.add_heading(doc_data["title"], level=0)
            
            # Write contents
            for block in doc_data["content"].split("\n\n"):
                if block.startswith("## "):
                    doc.add_heading(block.replace("## ", ""), level=1)
                else:
                    doc.add_paragraph(block)

            # References
            if doc_data["bibliography"]:
                doc.add_page_break()
                doc.add_heading("References", level=1)
                for ref in doc_data["bibliography"].split("\n\n")[1:]:
                    doc.add_paragraph(ref)
            
            doc.save(file_path)
            
        except ImportError:
            # Fallback to saving as TXT if python-docx isn't ready
            file_path = file_path.replace(".docx", ".txt")
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(f"{doc_data['title']}\n\n{doc_data['content']}\n\n{doc_data['bibliography']}")

        return file_path

    @staticmethod
    def export_as_html(db: Session, project_id: int, citation_style: str = "ieee") -> str:
        """
        Compiles the draft into a beautifully styled HTML file.
        """
        doc_data = ExportService.compile_document(db, project_id, citation_style)
        filename = f"Project_{project_id}_Draft.html"
        file_path = os.path.join(settings.DOCS_DIR, filename)

        # Convert markdown-like headers to HTML
        body_html = ""
        for block in doc_data["content"].split("\n\n"):
            if block.startswith("## "):
                sec = block.replace("## ", "")
                body_html += f"<h2>{sec}</h2>"
            else:
                body_html += f"<p>{block}</p>"

        ref_html = ""
        if doc_data["bibliography"]:
            ref_html += "<h2>References</h2>"
            for ref in doc_data["bibliography"].split("\n\n")[1:]:
                ref_html += f"<p class='reference'>{ref}</p>"

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{doc_data['title']}</title>
    <style>
        body {{
            font-family: 'Times New Roman', Times, serif;
            line-height: 2.0;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            color: #111;
        }}
        h1 {{ text-align: center; font-size: 24pt; margin-bottom: 30px; }}
        h2 {{ font-size: 14pt; margin-top: 30px; border-bottom: 1px solid #ccc; padding-bottom: 5px; }}
        p {{ font-size: 12pt; text-indent: 0.5in; margin: 0 0 15px 0; text-align: justify; }}
        .reference {{ font-size: 11pt; padding-left: 0.5in; text-indent: -0.5in; margin-bottom: 10px; }}
    </style>
</head>
<body>
    <h1>{doc_data['title']}</h1>
    {body_html}
    {ref_html}
</body>
</html>
"""
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        return file_path
